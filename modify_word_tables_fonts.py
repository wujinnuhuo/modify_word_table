import os
from docx import Document
from docx.shared import Pt


def format_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = 1  # Center alignment
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.bold = True
            run.font.size = Pt(12)  # You can adjust the font size as needed


def process_word_file(file_path, positions):
    doc = Document(file_path)

    if len(doc.tables) > 0:
        table = doc.tables[0]  # Only process the first table

        for row_idx, col_idx in positions:
            if 0 <= row_idx < len(table.rows) and 0 <= col_idx < len(table.columns):
                cell = table.cell(row_idx, col_idx)
                format_cell(cell)

        doc.save(file_path)  # Save changes directly to the source file


def process_folder(folder_path, positions):
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.endswith(".docx"):
                file_path = os.path.join(root, file)
                process_word_file(file_path, positions)


def main():
    folder_path = input("请输入文件夹的路径：")
    positions_input = input("请输入要修改的行列位置，以空格分隔，格式如：1,1 2,2 3,3：")
    positions = [tuple(map(int, pos.split(","))) for pos in positions_input.split()]

    process_folder(folder_path, positions)
    print("所有文件修改完成。")


if __name__ == "__main__":
    main()
